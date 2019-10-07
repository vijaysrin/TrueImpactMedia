import pyexcel as pe
import csv
from docx import Document
from docx.shared import Pt
from docx.shared import Inches
from datetime import date
from datetime import timedelta
from docx.enum.table import WD_TABLE_ALIGNMENT
#
def buckets(starts, days, ends):
	unique_days = {}
	for s in range(0, len(starts)):
		if(starts[s] in unique_days):
			unique_days[starts[s]].append(days[s])
		else:
			unique_days[starts[s]] = []
			unique_days[starts[s]].append(days[s])
	buckets = {}
	outliers = []
	for key in unique_days.keys():
		if(len(unique_days[key]) > 1):
			temp_list = unique_days[key]
			temp_list.sort()
			max_days = temp_list[-1]
			periods = max_days/28
			month_day_year = key.split('/')
			month = month_day_year[0]
			day = month_day_year[1]
			year = month_day_year[2]
			date2 = date(int(year), int(month), int(day))
			for i in range(0, int(periods)):
				str_date = str(date2.month) + "/" + str(date2.day) + "/" + str(date2.year)
				if(str_date not in buckets):
					buckets[str_date] = []
				date2 += timedelta(days=28)
		else:
			outliers.append(starts.index(key))
	return buckets, outliers

def add_boards(bucket_dict, starts, days, outliers_list):
	for a in range(0, len(starts)):
		if(a not in outliers_list):
			periods = days[a]/28
			month_day_year = starts[a].split('/')
			month = month_day_year[0]
			day = month_day_year[1]
			year = month_day_year[2]
			date2 = date(int(year), int(month), int(day))
			for i in range(0, int(periods)):
				str_date = str(date2.month) + "/" + str(date2.day) + "/" + str(date2.year)
				bucket_dict[str_date].append(a)
				date2 += timedelta(days=28)	
	return bucket_dict

def output_outliers(xlfile, outliers_list, data):
	xlfile = xlfile.replace(".csv", "")
	file_name = xlfile + "_outliers.csv"
	with open(file_name, 'w') as csv_file:
		writer = csv.writer(csv_file)

		outliers_list.sort()
		row1 = []
		for key in data.keys():
			row1.append(key)
		row1[0] = 'Purchase ID'
		writer.writerow(row1)
		for o in outliers_list:
			row = []
			for key in data.keys():
				row.append(data[key][o])
			writer.writerow(row)
		csv_file.close()
	print
	print('Outliers File created: ' + file_name)
	print

def create_invoices1(bucket_dict, operators, panels, rfps, starts, client_rate, print_cost):
	print(bucket_dict)
	template = Document('purchase_template.docx')

	for key in bucket_dict.keys():
		month_day_year = key.split("/")
		monthdayyear = month_day_year[0] + "-" + month_day_year[1] + "-" + month_day_year[2]

		month = month_day_year[0]
		day = month_day_year[1]
		year = month_day_year[2]
		date2 = date(int(year), int(month), int(day))
		date2 += timedelta(days=28)
		#if(date2 > )
		str_date = str(date2.month) + "/" + str(date2.day) + "/" + str(date2.year)
		total = 0
		for num in bucket_dict[key]:
			total += client_rate[num]
			if(starts[num] == key):
				total+= print_cost[num]

		#creates new file
		template.save(monthdayyear + '.docx')

		#write to new file
		document = Document(monthdayyear + ".docx")
		style = document.styles['Normal']
		font = style.font
		font.name = 'Arial'
		font.size = Pt(9)

		#first table
		document.tables[1].row_cells(1)[1].text = key
		document.tables[1].row_cells(1)[2].text = str_date
		document.tables[1].row_cells(1)[3].text = '${:,.2f}'.format(total)

		#second table
		rownum = 1
		boardcost = 0
		index = 0
		for n in bucket_dict[key]:
			document.tables[2].row_cells(rownum)[1].text = operators[n]
			document.tables[2].row_cells(rownum)[0].text = panels[n]
			document.tables[2].row_cells(rownum)[2].text = rfps[n]
			#document.tables[2].row_cells(rownum)[3].text = key
			document.tables[2].row_cells(rownum)[3].text = key + " - " + str_date
			boardcost = client_rate[n]
			print
			print(boardcost)
			if(starts[n] == key):
				boardcost += print_cost[n]
			print(panels[n] + "\n" + str(boardcost))
			print
			document.tables[2].row_cells(rownum)[4].text = '${:,.2f}'.format(boardcost)
			rownum+=1
			document.tables[2].add_row()
			index = n

		document.tables[2].row_cells(n+1)[4].text = '${:,.2f}'.format(total)
		document.tables[2].row_cells(n+2)[4].text = '${:,.2f}'.format(0)
		document.tables[2].row_cells(n+3)[4].text = '${:,.2f}'.format(total)

		#save file
		document.save(monthdayyear + ".docx")
#

def earliest_date(boards):
	dates = boards[0][3].split("/")
	earliest_date = date(int(dates[2]) + 2000, int(dates[0]), int(dates[1]))
	for i in range(1, boards.__len__()):
		new_date = boards[i][3].split("/")
		value = date(int(new_date[2]) + 2000, int(new_date[0]), int(new_date[1]))
		if(value < earliest_date):
			earliest_date = value
	return earliest_date

def latest_date(boards):
	dates = boards[0][4].split("/")
	latest_date = date(int(dates[2]) + 2000, int(dates[0]), int(dates[1]))
	for i in range(1, boards.__len__()):
		new_date = boards[i][4].split("/")
		value = date(int(new_date[2]) + 2000, int(new_date[0]), int(new_date[1]))
		if(value > latest_date):
			latest_date = value
	return latest_date

def date_toString(current_date):
	year = current_date.year
	month = current_date.month
	day = current_date.day
	result = str(month) + "/" + str(day) + "/" + str(year - 2000)
	return result

def date_toString2(current_date):
	year = current_date.year
	month = current_date.month
	day = current_date.day
	result = str(month) + "-" + str(day) + "-" + str(year - 2000)
	return result


def stringToDate(datex):
	dates = datex.split("/")
	new_date = date(int(dates[2]) + 2000, int(dates[0]), int(dates[1]))
	return new_date


def createInvoices(boards):
	template = Document('purchase_template.docx')
	current_start = earliest_date(boards)
	latest_day = latest_date(boards)
	print (latest_day)
	tag = False
	while(current_start < latest_day and not tag):
		#creates new file
		
		template.save(date_toString2(current_start) + '.docx')

		#write to new file
		document = Document(date_toString2(current_start) + ".docx")
		style = document.styles['Normal']
		font = style.font
		font.name = 'Arial'
		font.size = Pt(9)
		#document.table[1].alignment = WD_TABLE_ALIGNMENT.CENTER
		document.tables[1].row_cells(1)[1].text = date_toString(current_start)
		document.tables[1].row_cells(1)[2].text = date_toString(current_start + timedelta(days=28))
		
		rownum = 1
		totalCost = 0
		start_range = current_start
		end_range = start_range + timedelta(days=28)
		print("Here")
		for i in range(0, boards.__len__()):
			
			
			if(start_range <= stringToDate(boards[i][3]) <= end_range):
				print("INSIDDE IF")
				document.tables[2].row_cells(rownum)[0].text = boards[i][0]
				document.tables[2].row_cells(rownum)[1].text = boards[i][1]
				document.tables[2].row_cells(rownum)[2].text = boards[i][2]

				monthly = boards[i][9]
				if(stringToDate(boards[i][3]) == stringToDate(boards[i][10])):
					monthly += (boards[i][7] + boards[i][8])
				totalCost += monthly
				document.tables[2].row_cells(rownum)[4].text = '${:,.2f}'.format(monthly)

				if((stringToDate(boards[i][3]) + timedelta(days=28)) <= stringToDate(boards[i][4])):
					if(stringToDate(boards[i][4]) - (stringToDate(boards[i][3]) + timedelta(days=28)) < timedelta(days=20)):
						document.tables[2].row_cells(rownum)[3].text = boards[i][3] + " - " + boards[i][4]
						if(stringToDate(boards[i][4]) == latest_day ):
							tag = True
						#boards[i][3] = boards[i][4]
					else:
						document.tables[2].row_cells(rownum)[3].text = boards[i][3] + " - " + date_toString(stringToDate(boards[i][3]) + timedelta(days=28))
						if((stringToDate(boards[i][3]) + timedelta(days=28)) < stringToDate(boards[i][4])):
							boards[i][3] = date_toString(stringToDate(boards[i][3]) + timedelta(days=29))
				else:
					document.tables[2].row_cells(rownum)[3].text = boards[i][3] + " - " + boards[i][4]
					boards[i][3] = boards[i][4]
				if(i != (boards.__len__() - 1)):
					document.tables[2].add_row()
				rownum += 1
		document.tables[1].row_cells(1)[3].text = '${:,.2f}'.format(totalCost)
		document.tables[2].add_row()
		document.tables[2].row_cells(rownum)[3].text = "SubTotal:"
		#document.tables[2].row_cells(rownum)[3].text.bold = True
		document.tables[2].row_cells(rownum)[4].text = '${:,.2f}'.format(totalCost)
		document.tables[2].add_row()
		document.tables[2].row_cells(rownum+1)[3].text = "Tax:"
		#document.tables[2].row_cells(rownum+1)[3].text.bold = True
		document.tables[2].row_cells(rownum+1)[4].text = '${:,.2f}'.format(0)
		document.tables[2].add_row()
		document.tables[2].row_cells(rownum+2)[3].text = "Total Due:"
		#document.tables[2].row_cells(rownum+2)[3].text.bold = True
		document.tables[2].row_cells(rownum+2)[4].text = '${:,.2f}'.format(totalCost)


		document.save(date_toString2(current_start) + '.docx')
		current_start += timedelta(days=29)
		
		#print (latest_day)

		



def main():
	#get file name from user
	xlfile = ""
	try:
		
		xlfile = input("Name of exact csv purchase file: ")
	except NameError:
		print ('Name Error')
		pass
	#open csv file
	csvfile = open(xlfile)

	#define columns by row 0 names
	data = pe.get_dict(file_name=xlfile, name_columns_by_row=0)

	#get data and assign to variables using row 0 names
	operators = data['Operator']
	panels = data['TIM Panel ID']
	rfps = data['RFP']
	starts = data['Contracted Start']
	ends = data['Contracted End']
	days = data['Duration Days']
	client_rate = data['Client Total']
	print_cost = data['Client Print Cost']
	install = data['Client Install Cost']

	#boards = []
	boards = [["" for x in range(11)] for y in range(data['Operator'].__len__())]
	#print(install)
	for x in range(0, data['Operator'].__len__()):
		
		boards[x][0] = panels[x]
		boards[x][1] = operators[x]
		boards[x][2] = rfps[x]
		boards[x][3] = starts[x]
		boards[x][4] = ends[x]
		boards[x][5] = days[x]
		boards[x][6] = client_rate[x]
		boards[x][7] = print_cost[x]
		boards[x][8] = install[x]
		initial = int(install[x]) + int(print_cost[x])
		cost = client_rate[x] - initial
		periods = int(days[x] / 28)
		if(periods == 0):
			periods = 1
		rate = cost / periods
		boards[x][9] = rate
		boards[x][10] = starts[x]

		

	createInvoices(boards)

	#print (latest_date(boards))
	
	#create buckets and get dictionary back
	#bucket_dict, outliers_list = buckets(starts, days, ends)

	#full_buckets = add_boards(bucket_dict, starts, days, outliers_list)

	#output_outliers(xlfile, outliers_list, data)

	#create_invoices(bucket_dict, operators, panels, rfps, starts, client_rate, print_cost)




main()